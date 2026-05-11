import os
import re
import io
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def render_latex_to_image(latex_str):
    """Renders a LaTeX string to a BytesIO image buffer using Matplotlib's internal mathtext."""
    try:
        # Create a figure with a very small size initially
        fig = plt.figure(figsize=(0.1, 0.1), dpi=300)
        # Add text. We use basic math mode delimiters
        if not latex_str.startswith('$'):
            render_str = f"${latex_str}$"
        else:
            render_str = latex_str
            
        text = fig.text(0, 0, render_str, fontsize=14, va='bottom', ha='left')
        
        # Draw to get bbox
        renderer = fig.canvas.get_renderer()
        bbox = text.get_window_extent(renderer=renderer)
        
        # Resize figure to fit text
        bbox_inches = bbox.transformed(fig.dpi_scale_trans.inverted())
        fig.set_size_inches(bbox_inches.width + 0.1, bbox_inches.height + 0.1)
        
        # Save to buffer
        buf = io.BytesIO()
        fig.savefig(buf, format='png', bbox_inches='tight', pad_inches=0.02, transparent=True)
        buf.seek(0)
        plt.close(fig)
        return buf
    except Exception as e:
        print(f"Error rendering latex '{latex_str}': {e}")
        return None

def set_table_borders(table):
    """Force grid borders on a table."""
    tbl = table._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        tcBorders = OxmlElement('w:tcBorders')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '4')
        tcBorders.append(top)
        
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        tcBorders.append(bottom)
        
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '4')
        tcBorders.append(left)
        
        right = OxmlElement('w:right')
        right.set(qn('w:val'), 'single')
        right.set(qn('w:sz'), '4')
        tcBorders.append(right)
        
        tcPr.append(tcBorders)

def convert_md_to_docx(md_path, docx_path, figure_dir):
    print(f"Starting conversion: {md_path} -> {docx_path}")
    
    doc = Document()
    
    # Set default styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Ensure headings also use Times New Roman
    for i in range(4):
        style_name = f'Heading {i}' if i > 0 else 'Title'
        if style_name in doc.styles:
            h_style = doc.styles[style_name]
            h_style.font.name = 'Times New Roman'
            h_style.font.bold = True
            h_style.font.color.rgb = RGBColor(0, 0, 0)
            if i == 0:
                h_style.font.size = Pt(16)
            else:
                h_style.font.size = Pt(14 - i)

    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    current_table = None
    in_code_block = False
    
    for line in lines:
        line_clean = line.strip()
        
        # Handle Code Blocks (for Algorithms)
        if line_clean.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph(line_clean)
            if 'No Spacing' in doc.styles:
                p.style = doc.styles['No Spacing']
            run = p.runs[0] if p.runs else p.add_run(line_clean)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            continue

        # Handle Tables (Markdown style)
        if '|' in line_clean and line_clean.count('|') >= 2:
            if '---' in line_clean:
                continue # Skip the markdown separator line
                
            parts = [p.strip() for p in line_clean.split('|')]
            if not parts[0]: parts = parts[1:]
            if not parts[-1]: parts = parts[:-1]
            
            if not parts:
                continue

            # Start of a new table or continuing one
            if current_table is None:
                current_table = doc.add_table(rows=0, cols=len(parts))
                current_table.style = 'Table Grid'
            
            row = current_table.add_row()
            for i, text in enumerate(parts):
                if i < len(row.cells):
                    row.cells[i].text = text.replace('**', '') 
            continue 
        else:
            if current_table:
                # set_table_borders(current_table) # Optional, Table Grid usually suffices
                current_table = None

        # Handle Headings
        if line_clean.startswith('# '):
            p = doc.add_heading(line_clean[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line_clean.startswith('## '):
            doc.add_heading(line_clean[3:], level=1)
        elif line_clean.startswith('### '):
            doc.add_heading(line_clean[4:], level=2)
            
        # Handle Figures (Support all 10)
        elif "Figure" in line_clean and (".png" in line_clean or "referencing" in line_clean):
            # Extract filename
            match = re.search(r'figure\d+_\w+\.png', line_clean)
            if match:
                fig_filename = match.group(0)
                # Check if file exists
                img_path = os.path.join(figure_dir, fig_filename)
                
                # Add caption first (Springer usually puts caption below, but we can stick to standard)
                # Actually, standard is Figure then Caption below.
                if os.path.exists(img_path):
                    p_fig = doc.add_paragraph()
                    p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p_fig.add_run()
                    # Resize to 3.2 inches (approx half width of page or small enough)
                    run.add_picture(img_path, width=Inches(3.2))
                    
                    # Add caption paragraph
                    p_cap = doc.add_paragraph(line_clean, style='Caption')
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    print(f"Warning: {img_path} not found.")
                    doc.add_paragraph(line_clean, style='Caption')
            else:
                 doc.add_paragraph(line_clean, style='Caption')
            continue
            
        # Handle Math Equations ($$ ... $$)
        elif line_clean.startswith('$$') and line_clean.endswith('$$'):
            latex_eq = line_clean[2:-2].strip()
            img_buf = render_latex_to_image(latex_eq)
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if img_buf:
                # Add the rendered equation image
                r = p.add_run()
                r.add_picture(img_buf)
            else:
                # Fallback to text if rendering fails
                r = p.add_run(latex_eq)
                r.font.italic = True
                r.font.color.rgb = RGBColor(255, 0, 0) # Red to indicate error
            continue

        # Handle Paragraphs covering text and inline formatting
        elif line_clean:
            if not line_clean.startswith('!['):
                p = doc.add_paragraph()
                
                # Handling bold conversions
                parts = re.split(r'(\*\*.*?\*\*)', line_clean)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
        else:
            pass

    doc.save(docx_path)
    print(f"Success! Saved to {docx_path}")

if __name__ == "__main__":
    import sys
    BASE_DIR = r"c:\Users\Lenovo\OneDrive\Desktop\6TH SEM\case study\case study"
    FIGURE_DIR = os.path.join(BASE_DIR, "_final_submission_", "project_figures")
    
    if len(sys.argv) > 2:
        MD_INPUT = sys.argv[1]
        DOCX_OUTPUT = sys.argv[2]
    else:
        # Default to Introduction when running standalone for test
        MD_INPUT = os.path.join(BASE_DIR, "springer_paper", "Introduction.md")
        DOCX_OUTPUT = os.path.join(BASE_DIR, "springer_paper", "Introduction.docx")
    
    convert_md_to_docx(MD_INPUT, DOCX_OUTPUT, FIGURE_DIR)
