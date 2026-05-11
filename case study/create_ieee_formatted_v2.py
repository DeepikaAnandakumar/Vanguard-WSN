"""
Complete IEEE-formatted Word document generator v2
Preserves ALL original content, includes actual images and tables
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_two_column_section(section):
    """Set section to two-column layout"""
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0] if sectPr.xpath('./w:cols') else OxmlElement('w:cols')
    if not sectPr.xpath('./w:cols'):
        sectPr.append(cols)
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '720')

def add_para(doc, text, bold=False, italic=False, size=10, align='justify', keep_with_next=False):
    """Helper to add formatted paragraph"""
    para = doc.add_paragraph(text)
    if align == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify':
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == 'left':
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    if keep_with_next:
        para.paragraph_format.keep_with_next = True
    
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return para

def add_image(doc, image_name, caption):
    """Add image with caption"""
    # Image path
    img_path = os.path.join('extracted_vanguard_content', 'word', 'media', image_name)
    if os.path.exists(img_path):
        # Add image, width constrained to column width approx
        try:
            doc.add_picture(img_path, width=Inches(3.4))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"Error adding image {image_name}: {e}")
            add_para(doc, f"[Image {image_name} missing]", italic=True, align='center')
    else:
        add_para(doc, f"[Image file {image_name} not found]", italic=True, align='center')
    
    # Add caption
    add_para(doc, caption, italic=True, size=9, align='center')
    doc.add_paragraph() # Spacing

def add_table(doc, data, caption):
    """Add formatted table"""
    add_para(doc, caption, bold=True, size=9, align='center', keep_with_next=True)
    
    rows = len(data)
    cols = len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.autofit = False 
    
    for r in range(rows):
        row_cells = table.rows[r].cells
        for c in range(cols):
            cell = row_cells[c]
            cell.text = str(data[r][c])
            # Format cell text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(8)
                    if r == 0: # Header row
                        run.bold = True
    
    doc.add_paragraph() # Spacing

# Load extracted content
with open('vanguard_content.txt', 'r', encoding='utf-8') as f:
    content_lines = [line.strip() for line in f.readlines()]

doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)

# --- HEADER SECTION (Single Column) ---
add_para(doc, content_lines[0], bold=True, size=24, align='center') # Title
doc.add_paragraph()
add_para(doc, content_lines[2], size=10, align='center') # Authors
add_para(doc, content_lines[3], italic=True, size=10, align='center') # Affiliation
add_para(doc, f"{content_lines[4]}, {content_lines[5]}", size=9, align='center') # Emails
doc.add_paragraph()

# Abstract
add_para(doc, "Abstract", bold=True, size=10, align='left')
add_para(doc, content_lines[8], italic=True, size=9, align='justify') # Abstract text
doc.add_paragraph()

# Keywords
add_para(doc, "Keywords", bold=True, italic=True, size=10, align='left')
add_para(doc, content_lines[9].replace("Keywords: ", ""), italic=True, size=9, align='left')
doc.add_paragraph()

# --- CONTENT SECTION (Two Columns) ---
set_two_column_section(doc.sections[0])

# Helper to skip lines we've processed manually
skip_lines = list(range(0, 11)) # Header + Abstract + Keywords

# Process body
for i, line in enumerate(content_lines):
    if i in skip_lines or not line:
        continue
    
    # Check for specific headings to insert tables/figures
    
    # Table I insertion (around Section 3.4)
    if "The table below places Vanguard-WSN" in line:
        add_para(doc, line, size=10)
        # Insert Table I
        data_t1 = [
            ["Feature Classification", "LEACH", "HEED", "PEGASIS", "Vanguard (Proposed)"],
            ["Primary Objective", "Load Rotation", "Energy balance", "Distance Min", "Utility Max"],
            ["Selection Logic", "Probabilistic", "Iterative Prob.", "Greedy Neighbor", "Deterministic"],
            ["Topology", "Star", "Cluster", "Chain", "EBPT Tree"],
            ["Heterogeneity", "None", "Initial Energy", "None", "Adaptive"],
            ["Cost", "Low", "High", "Medium", "Low (O(N))"],
            ["Recovery", "None", "Re-cluster", "Re-build", "Self-Healing"],
            ["Scalability", "Poor", "Medium", "Poor", "High"]
        ]
        add_table(doc, data_t1, "Table I: Qualitative Comparison")
        continue

    # Figure 1
    if "Figure 1:" in line:
        add_image(doc, "image1.png", line)
        continue
    
    # Figure 2
    if "Figure 2:" in line:
        add_image(doc, "image2.png", line)
        continue
        
    # Figure 3
    if "Figure 3:" in line:
        add_image(doc, "image3.png", line)
        continue

    # Table II (Table 1 in text)
    if "Table 1:" in line:
        # Insert Table II
        data_t2 = [
            ["Parameter", "Value"],
            ["Network Size", "100 Nodes"],
            ["Area", "100m x 100m"],
            ["BS Location", "(50, 50)"],
            ["Initial Energy", "0.5 J"],
            ["Data Pkt Size", "4000 bits"],
            ["Ctrl Pkt Size", "200 bits"]
        ]
        add_table(doc, data_t2, "Table II: Simulation Settings and Parameters")
        continue

    # Figure 4
    if "Figure 4:" in line:
        add_image(doc, "image4.png", line)
        continue

    # Figure 5
    if "Figure 5:" in line:
        add_image(doc, "image5.png", line)
        continue

    # Figure 6
    if "Figure 6:" in line:
        add_image(doc, "image6.png", line)
        continue

    # Figure 7 and 9 (Special case, appearing near each other)
    if "Figure 7:" in line:
        # Handles "Figure 7: ... Figure 9: ..." on same line
        parts = line.split("Figure 9:")
        caption7 = parts[0].strip()
        add_image(doc, "image7.png", caption7)
        
        if len(parts) > 1:
            caption9 = "Figure 9:" + parts[1].strip()
            add_image(doc, "image8.png", caption9) # Image 8 is Figure 9 based on order
        continue
    
    # Figure 8 (Heatmap, appearing after 7/9 in text)
    if "Figure 8:" in line:
        add_image(doc, "image9.png", line) # Image 9 is Figure 8 based on order
        continue

    # Table III (Table 2 in text)
    if "Table 2:" in line:
        # Insert Table III
        data_t3 = [
            ["Protocol", "FND (Rnd)", "Throughput", "Efficiency"],
            ["LEACH", "97.3", "1x", "Low"],
            ["HEED", "210", "2.1x", "Med"],
            ["PEGASIS", "350", "1.5x", "High"],
            ["Vanguard", "993.1", "10.2x", "V. High"]
        ]
        add_table(doc, data_t3, "Table III: Comparative Performance Metrics")
        continue

    # Standard Paragraph Processing
    if line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ')):
        add_para(doc, line, bold=True, size=12, align='left')
    elif any(line.startswith(f'{j}.{k}') for j in range(1, 7) for k in range(1, 20)):
        add_para(doc, line, bold=True, size=10, align='left')
    elif any(line.startswith(f'{j}.{k}.{m}') for j in range(1, 7) for k in range(1, 20) for m in range(1, 10)):
        add_para(doc, line, italic=True, size=10, align='left')
    elif line == 'References':
        add_para(doc, line, bold=True, size=12, align='left')
    elif line.startswith('['):
        add_para(doc, line, size=8, align='left')
    else:
        add_para(doc, line, size=10, align='justify')

doc.save('Vanguard_IEEE_Formatted.docx')
print("Vanguard_IEEE_Formatted.docx (v2) created with images and tables.")
