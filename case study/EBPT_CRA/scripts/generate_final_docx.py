import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx():
    # Paths
    base_dir = r"c:\Users\Lenovo\OneDrive\Desktop\6TH SEM\case study\case study\EBPT_CRA"
    md_path = os.path.join(base_dir, "submission_paper", "Vanguard_WSN_IEEE_Paper.md")
    output_path = os.path.join(base_dir, "submission_paper", "Vanguard_WSN_Final_v3.docx")
    figs_dir = os.path.join(base_dir, "project_figures")
    
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found")
        return

    doc = Document()
    
    # Set default font to Times New Roman if possible
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    print("Parsing Markdown and building Word document...")
    
    in_table = False
    table_data = []

    for line in lines:
        line = line.strip()
        
        # Handle Headings
        if line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=0)
            continue
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=1)
            continue
        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=2)
            continue
            
        # Handle Images (Markdown syntax: ![caption](path))
        img_match = re.search(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            caption = img_match.group(1)
            img_rel_path = img_match.group(2)
            # Use absolute path
            img_abs_path = os.path.join(base_dir, img_rel_path.replace("/", "\\"))
            
            if os.path.exists(img_abs_path):
                doc.add_picture(img_abs_path, width=Inches(5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add caption below
                p = doc.add_paragraph(f"Figure: {caption}")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.style = doc.styles['Caption'] if 'Caption' in doc.styles else 'Normal'
            else:
                print(f"Warning: Image not found at {img_abs_path}")
            continue

        # Handle Tables (Markdown syntax: | col1 | col2 |)
        if line.startswith("|") and "---" not in line:
            in_table = True
            cells = [c.strip() for c in line.split("|") if c.strip()]
            table_data.append(cells)
            continue
        elif in_table and (not line.startswith("|") or "---" in line):
            if "---" in line:
                continue
            # End of table, render it
            if table_data:
                rows = len(table_data)
                cols = len(table_data[0]) if rows > 0 else 0
                if cols > 0:
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid'
                    for r in range(rows):
                        for c in range(cols):
                            # Handle cases where row might have fewer cells
                            if c < len(table_data[r]):
                                table.cell(r, c).text = table_data[r][c]
                table_data = []
            in_table = False
            if not line: continue

        # Handle Lists
        if line.startswith("* ") or line.startswith("- "):
            doc.add_paragraph(line[2:], style='List Bullet')
            continue
        elif re.match(r'^\d+\.\s', line):
            doc.add_paragraph(line[3:], style='List Number')
            continue

        # Regular Text (with simple bold/italic)
        if line and not in_table:
            # Replace markdown bold/italic with docx formatting is complex, 
            # for now just add paragraph. We'll improve bolding for key terms.
            para = doc.add_paragraph()
            
            # Simple bold/italic regex replacement
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    para.add_run(part[2:-2]).bold = True
                elif part.startswith("*") and part.endswith("*"):
                    para.add_run(part[1:-1]).italic = True
                else:
                    para.add_run(part)

    # Append all remaining Figures 1-10 at the end as a "Results Appendix" if not inserted
    doc.add_page_break()
    doc.add_heading("PROCESSED RESULTS AND FIGURES", level=1)
    
    processed_figures = [
        ("figure1_architecture.png", "Figure 1: Network Architecture Schematic"),
        ("figure2_deployment.png", "Figure 2: Node Deployment and Sink Localization"),
        ("figure3_routing_tree.png", "Figure 3: Energy-Balanced Path Tree (EBPT) Construction"),
        ("figure4_pareto.png", "Figure 4: Multi-Objective Optimization (Lifetime vs. Fairness)"),
        ("figure5_lifetime.png", "Figure 5: Network Lifetime Survival Curves (Alive Nodes)"),
        ("figure6_death_rounds.png", "Figure 6: Significant Death Epochs (FND, HND, LND)"),
        ("figure7_fairness_monitoring.png", "Figure 7: Jain's Fairness Index stability Profile"),
        ("figure8_throughput.png", "Figure 8: Cumulative Throughput with Shaded Error Bands"),
        ("figure9_heatmap.png", "Figure 9: Spatial Energy Depletion Heatmap (Exact Results)"),
        ("figure10_snapshot.png", "Figure 10: Final Network Topology Snapshot at Round 500")
    ]
    
    for filename, caption in processed_figures:
        path = os.path.join(figs_dir, filename)
        if os.path.exists(path):
            doc.add_heading(caption, level=2)
            doc.add_picture(path, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(caption).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph() # Spacer

    doc.save(output_path)
    print(f"Success! Document saved to {output_path}")

if __name__ == "__main__":
    generate_docx()
