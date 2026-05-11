"""
Complete IEEE-formatted Word document generator for Vanguard WSN paper
Preserves ALL original content with proper IEEE conference formatting
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_two_column_section(section):
    """Set section to two-column layout"""
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0] if sectPr.xpath('./w:cols') else OxmlElement('w:cols')
    if not sectPr.xpath('./w:cols'):
        sectPr.append(cols)
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '720')

def add_para(doc, text, bold=False, italic=False, size=10, align='justify'):
    """Helper to add formatted paragraph"""
    para = doc.add_paragraph(text)
    if align == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify':
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return para

# Read original content
with open('vanguard_content.txt', 'r', encoding='utf-8') as f:
    content_lines = f.readlines()

# Create document
doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)

# TITLE (from line 1)
add_para(doc, content_lines[0].strip(), bold=True, size=24, align='center')

doc.add_paragraph()

# AUTHORS (from lines 3)
add_para(doc, content_lines[2].strip(), size=10, align='center')

# AFFILIATION (from line 4)
add_para(doc, content_lines[3].strip(), italic=True, size=10, align='center')

# EMAILS (from lines 5-6)
email_text = content_lines[4].strip() + ", " + content_lines[5].strip()
add_para(doc, email_text, size=9, align='center')

doc.add_paragraph()

# ABSTRACT HEADING
add_para(doc, "Abstract", bold=True, size=10, align='left')

# ABSTRACT TEXT (from line 9)
add_para(doc, content_lines[8].strip(), italic=True, size=10, align='justify')

# KEYWORDS (from line 10)
add_para(doc, "Keywords", bold=True, italic=True, size=10, align='left')
add_para(doc, content_lines[9].strip().replace("Keywords: ", ""), italic=True, size=10, align='left')

# Switch to two columns
set_two_column_section(doc.sections[0])

# Parse and add all remaining content
i = 10
while i < len(content_lines):
    line = content_lines[i].strip()
    
    if not line:
        i += 1
        continue
    
    # Main section headings (1., 2., 3., etc.)
    if line.startswith(('1. Introduction', '2. Main Contribution', '3. Background', 
                        '4. Methodology', '5. Simulation', '6. Conclusion')):
        add_para(doc, line, bold=True, size=12, align='left')
    
    # Subsection headings (1.1, 1.2, etc.)
    elif any(line.startswith(f'{j}.{k}') for j in range(1, 7) for k in range(1, 20)):
        add_para(doc, line, bold=True, size=10, align='left')
    
    # Sub-subsection headings (1.1.1, 1.2.1, etc.)
    elif any(line.startswith(f'{j}.{k}.{m}') for j in range(1, 7) for k in range(1, 20) for m in range(1, 10)):
        add_para(doc, line, italic=True, size=10, align='left')
    
    # Figures
    elif line.startswith('Figure'):
        add_para(doc, f"[{line}]", italic=True, size=9, align='center')
    
    # Tables
    elif line.startswith('Table'):
        add_para(doc, line, bold=True, size=9, align='center')
    
    # References heading
    elif line == 'References':
        add_para(doc, line, bold=True, size=12, align='left')
    
    # Reference items
    elif line.startswith('['):
        add_para(doc, line, size=8, align='left')
    
    # Algorithm
    elif line.startswith('Algorithm'):
        add_para(doc, line, bold=True, size=10, align='left')
    
    # Regular text
    else:
        add_para(doc, line, size=10, align='justify')
    
    i += 1

# Save
doc.save('Vanguard_IEEE_Formatted.docx')
print("✓ IEEE Word document created: Vanguard_IEEE_Formatted.docx")
