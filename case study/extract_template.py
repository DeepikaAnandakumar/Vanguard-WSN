from docx import Document
import sys

def extract_template_info(docx_path, output_path):
    """Extract template structure and formatting info from a Word document"""
    try:
        doc = Document(docx_path)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"=== Template Analysis: {docx_path} ===\n\n")
            
            # Extract paragraphs with style information
            f.write("--- PARAGRAPH STYLES ---\n")
            for i, para in enumerate(doc.paragraphs[:50]):  # First 50 paragraphs
                if para.text.strip():
                    f.write(f"\nPara {i+1}:\n")
                    f.write(f"  Style: {para.style.name}\n")
                    f.write(f"  Text: {para.text[:100]}\n")
            
            # Extract section information
            f.write("\n\n--- SECTIONS ---\n")
            for i, section in enumerate(doc.sections):
                f.write(f"\nSection {i+1}:\n")
                f.write(f"  Page width: {section.page_width}\n")
                f.write(f"  Page height: {section.page_height}\n")
                f.write(f"  Top margin: {section.top_margin}\n")
                f.write(f"  Bottom margin: {section.bottom_margin}\n")
                f.write(f"  Left margin: {section.left_margin}\n")
                f.write(f"  Right margin: {section.right_margin}\n")
            
            # List all styles used
            f.write("\n\n--- STYLES USED ---\n")
            styles_used = set()
            for para in doc.paragraphs:
                if para.text.strip():
                    styles_used.add(para.style.name)
            for style in sorted(styles_used):
                f.write(f"  - {style}\n")
        
        print(f"Success: {output_path}")
    except Exception as e:
        print(f"Error processing {docx_path}: {e}")
        sys.exit(1)

if __name__ == "__main__":
    # Process both templates
    extract_template_info('conference-template-a4-ICCSP.docx', 'template_iccsp_analysis.txt')
    extract_template_info('conference-template-a4.docx', 'template_a4_analysis.txt')
